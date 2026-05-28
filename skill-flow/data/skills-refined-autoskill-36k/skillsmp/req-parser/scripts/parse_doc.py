#!/usr/bin/env python3
"""
文档解析工具（改进版，支持跨平台和多种降级方案）

功能：
- 扫描输入目录下的 PDF/DOCX/DOC/PPTX
- 使用 Docling 转换为 Markdown（优先）
- 提取图片到 assets 目录
- 降级方案：PDF用PyPDF2，DOCX用python-docx
- 输出到指定目录

用法：
    uv run skills/req-parser/scripts/parse_doc.py --input-dir original-requirements --output-dir cleaned-requirements/chunks
"""

import argparse
import base64
import sys
from pathlib import Path

# 尝试导入核心库
docling_available = False
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.datamodel.base_models import InputFormat
    from docling.document_converter import PdfFormatOption, WordFormatOption
    docling_available = True
except ImportError:
    pass

pypdf2_available = False
try:
    import PyPDF2
    pypdf2_available = True
except ImportError:
    pass

python_docx_available = False
try:
    from docx import Document
    python_docx_available = True
except ImportError:
    pass


def convert_with_docling(file_path: Path, assets_dir: Path = None):
    """使用 Docling 转换文档，返回 (markdown_content, image_count)"""
    if not docling_available:
        return None, 0
    try:
        print(f"🚀 [Docling] 正在解析: {file_path.name}")

        # 配置 PDF Pipeline 选项
        pipeline_options = PdfPipelineOptions()
        pipeline_options.generate_picture_images = True
        pipeline_options.generate_table_images = True
        pipeline_options.images_scale = 2.0

        # 创建格式选项（仅 PDF 需要特殊配置，DOCX 使用默认）
        format_options = {
            InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
        }

        converter = DocumentConverter(format_options=format_options)
        result = converter.convert(file_path)

        # 提取并保存图片
        image_count = 0
        image_filenames = []
        if assets_dir and hasattr(result.document, 'pictures') and result.document.pictures:
            assets_dir.mkdir(parents=True, exist_ok=True)
            print(f"📷 发现 {len(result.document.pictures)} 张图片")

            for i, pic in enumerate(result.document.pictures):
                if hasattr(pic, 'image') and pic.image and hasattr(pic.image, 'uri'):
                    uri = str(pic.image.uri)

                    # 解析 base64 数据
                    if uri.startswith('data:image/'):
                        parts = uri.split(',', 1)
                        if len(parts) == 2:
                            mime_type = parts[0].split(':')[1].split(';')[0]
                            ext = mime_type.split('/')[-1]
                            base64_data = parts[1]

                            # 解码并保存
                            try:
                                image_data = base64.b64decode(base64_data)
                                image_filename = f"figure-{i+1}.{ext}"
                                image_path = assets_dir / image_filename
                                image_path.write_bytes(image_data)
                                image_filenames.append(image_filename)
                                image_count += 1
                                print(f"  ✅ 保存图片 {i+1}: {image_filename} ({len(image_data)} bytes)")
                            except Exception as img_err:
                                print(f"  ⚠️  图片 {i+1} 保存失败: {img_err}")
                                image_filenames.append(None)
                    else:
                        image_filenames.append(None)
                else:
                    image_filenames.append(None)

        markdown_content = result.document.export_to_markdown()

        # 替换占位符为带文件名的格式
        if image_filenames:
            lines = markdown_content.split('\n')
            new_lines = []
            image_index = 0
            for line in lines:
                if line.strip() == '<!-- image -->':
                    if image_index < len(image_filenames) and image_filenames[image_index]:
                        new_lines.append(f"<!-- image: {image_filenames[image_index]} -->")
                    else:
                        new_lines.append(line)
                    image_index += 1
                else:
                    new_lines.append(line)
            markdown_content = '\n'.join(new_lines)

        return markdown_content, image_count
    except Exception as e:
        print(f"⚠️  [Docling] 失败: {e}")
        return None, 0


def convert_pdf_with_pypdf2(file_path: Path) -> str:
    """使用 PyPDF2 转换 PDF"""
    if not pypdf2_available:
        return None
    try:
        print(f"⚠️  [PyPDF2] 降级解析 PDF: {file_path.name}")
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            num_pages = len(reader.pages)
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"## Page {page_num + 1}\n\n{text}\n")
        return "\n".join(text_parts) if text_parts else None
    except Exception as e:
        print(f"❌ [PyPDF2] 失败: {e}")
        return None


def convert_table_to_markdown(table) -> str:
    """将 Word 表格转换为 Markdown 表格"""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append('| ' + ' | '.join(cells) + ' |')
        if i == 0:
            rows.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
    return '\n'.join(rows)


def convert_docx_with_python_docx(file_path: Path) -> str:
    """使用 python-docx 转换 DOCX"""
    if not python_docx_available:
        return None
    try:
        print(f"⚠️  [python-docx] 降级解析 DOCX: {file_path.name}")
        doc = Document(file_path)
        text_parts = []

        # 提取段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    if level.isdigit():
                        text_parts.append(f"{'#' * int(level)} {text}\n")
                    else:
                        text_parts.append(f"## {text}\n")
                else:
                    text_parts.append(f"{text}\n")

        # 提取表格
        for table in doc.tables:
            md_table = convert_table_to_markdown(table)
            text_parts.append(f"\n{md_table}\n")

        return "\n".join(text_parts) if text_parts else None
    except Exception as e:
        print(f"❌ [python-docx] 失败: {e}")
        return None


def process_file(input_file: Path, output_dir: Path, assets_dir: Path):
    """处理单个文件"""
    output_path = output_dir / f"{input_file.stem}.md"
    if output_path.exists():
        print(f"⏩ 已存在，跳过: {output_path.name}")
        return

    # 优先使用 Docling
    content, image_count = convert_with_docling(input_file, assets_dir)

    # 根据文件类型选择降级方案
    if not content:
        suffix = input_file.suffix.lower()
        if suffix == '.pdf':
            content = convert_pdf_with_pypdf2(input_file)
        elif suffix in {'.docx', '.doc'}:
            content = convert_docx_with_python_docx(input_file)

    if content:
        output_path.write_text(content, encoding='utf-8')
        if image_count > 0:
            print(f"✅ 已写入: {output_path.name} (包含 {image_count} 张图片)")
        else:
            print(f"✅ 已写入: {output_path.name}")
    else:
        print(f"❌ 解析失败，无法提取内容: {input_file.name}")
        print(f"   提示: 请确保已安装依赖 (uv sync)")


def main():
    parser = argparse.ArgumentParser(description="文档解析器")
    parser.add_argument("--input-dir", type=Path, required=True, help="输入目录")
    parser.add_argument("--output-dir", type=Path, required=True, help="输出目录")
    args = parser.parse_args()

    if not args.input_dir.exists():
        print(f"❌ 输入目录不存在: {args.input_dir}")
        sys.exit(1)

    # 检查依赖
    if not (docling_available or pypdf2_available or python_docx_available):
        print("❌ 错误：未安装任何文档解析库")
        print("   请执行: uv sync")
        sys.exit(1)

    args.output_dir.mkdir(parents=True, exist_ok=True)

    # 创建 assets 目录（在 output_dir 的父目录下）
    assets_dir = args.output_dir.parent / "assets"

    supported_exts = {'.pdf', '.docx', '.doc', '.pptx'}
    found = False

    for f in args.input_dir.iterdir():
        if f.is_file() and f.suffix.lower() in supported_exts:
            found = True
            process_file(f, args.output_dir, assets_dir)
        elif f.is_file() and f.suffix.lower() in {'.md', '.txt'}:
            # 直接复制纯文本
            dest = args.output_dir / f"{f.stem}.md"
            dest.write_text(f.read_text(encoding='utf-8'), encoding='utf-8')
            print(f"📋 已复制文本: {f.name}")
            found = True

    if not found:
        print("ℹ️  未发现支持的文档 (PDF/DOCX/MD/TXT)")


if __name__ == "__main__":
    main()
