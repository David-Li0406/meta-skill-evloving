#!/usr/bin/env python3
"""
Extract Text from DOCX using python-docx

Usage:
    python3 extract_docx.py <docx_path>

Dependencies:
    pip install python-docx
"""

import sys
import argparse
import docx


def extract_text(docx_path):
    """
    Extracts full text from a DOCX file.

    Args:
        docx_path (str): Path to the .docx file.

    Returns:
        str: extracted text
    """
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"Error opening DOCX: {e}", file=sys.stderr)
        sys.exit(1)

    full_text = []

    # Paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Tables (optional, but good for resumes)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))

    return "\n".join(full_text)


def main():
    parser = argparse.ArgumentParser(description="Extract text from DOCX")
    parser.add_argument("docx_path", help="Path to the DOCX file")

    args = parser.parse_args()

    text = extract_text(args.docx_path)
    print(text)


if __name__ == "__main__":
    main()
